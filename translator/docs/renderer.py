# -*- coding: utf-8 -*-
"""
renderer.py — 文档重建与导出模块
将翻译后的 DocumentModel 渲染为:
- .docx: python-docx 重建(保留段落样式/字体/加粗/表格/列表)
- .pdf : reportlab 生成(标题/段落/列表/表格/分页, 自动注册中文字体)

导出格式说明:
- Word 翻译后导出 .docx: 结构级版式完整保留(标题层级、粗体、字号、对齐、表格)
- PDF 翻译后导出 .pdf: 由于 PDF 无结构化版式, 采用"版式重建"策略
  (按内容结构重新排版), 保证内容完整与排版美观, 非逐像素还原。
"""
import os

from .block import DocumentModel, KIND_HEADING, KIND_LIST_ITEM, KIND_EMPTY

# 中文粗体字体(Windows 微软雅黑; 按顺序探测)
_CJK_BOLD_FONTS = [
    "C:/Windows/Fonts/msyhbd.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc",
]
_CJK_REGULAR_FONTS = [
    "C:/Windows/Fonts/msyh.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
]


# ---------------------------------------------------------------
# Word (.docx) 导出
# ---------------------------------------------------------------
def render_docx(model: DocumentModel, out_path: str):
    """将文档模型渲染为 .docx"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()

    # 文档默认字体(兼顾中文)
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(11)
    except Exception:
        pass

    _apply_align = {
        0: WD_ALIGN_PARAGRAPH.LEFT,
        1: WD_ALIGN_PARAGRAPH.CENTER,
        2: WD_ALIGN_PARAGRAPH.RIGHT,
        3: WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    def add_paragraph(text: str, kind: str, style: dict, meta: dict):
        if kind == KIND_EMPTY or not text.strip():
            doc.add_paragraph("")
            return
        p = doc.add_paragraph()
        run = p.add_run(text)

        if style.get("bold"):
            run.bold = True
        if style.get("italic"):
            run.italic = True
        if style.get("font_size"):
            run.font.size = Pt(style["font_size"])
        align = style.get("align")
        if align in _apply_align:
            p.alignment = _apply_align[align]

        if kind == KIND_HEADING:
            level = meta.get("level", 1)
            p.style = doc.styles[f"Heading {min(level, 4)}"]
            run.bold = True
        elif kind == KIND_LIST_ITEM:
            p.style = doc.styles["List Bullet"]
        return p

    # 顺序块
    for b in model.blocks:
        add_paragraph(b.text, b.kind, b.style, b.meta)

    # 表格
    for table_model in model.tables:
        if not table_model:
            continue
        n_rows = len(table_model)
        n_cols = max(len(r) for r in table_model)
        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.style = "Table Grid"
        for r, row in enumerate(table_model):
            for c, cell_blocks in enumerate(row):
                if c >= n_cols:
                    continue
                cell = table.cell(r, c)
                # 清空默认段落, 写入单元格块
                cell.text = ""
                for cb in cell_blocks:
                    para = cell.paragraphs[0] if cb is cell_blocks[0] and not cell.paragraphs[0].text else cell.add_paragraph()
                    para.add_run(cb.text)
        doc.add_paragraph("")  # 表格后空行

    doc.save(out_path)


# ---------------------------------------------------------------
# PDF 导出
# ---------------------------------------------------------------
def _register_cjk_fonts() -> tuple:
    """
    注册中文字体到 reportlab, 返回 (普通字体名, 粗体字体名)。

    优先级:
    1. 外部 TTF/TTC 字体(逐个尝试, 注册失败自动跳过——
       例如 NotoSansCJK.ttc 的 PostScript 轮廓 reportlab 不支持, 会抛异常)
    2. reportlab 内置 CID 字体 STSong-Light(零依赖, Docker/任意环境可用, 支持简体中文)
    3. Helvetica(最终兜底, 中文将无法显示)
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    registered = pdfmetrics.getRegisteredFontNames()

    # 1. 尝试外部字体文件(Windows msyh / Linux Noto)
    for reg_path in _CJK_REGULAR_FONTS:
        if not os.path.exists(reg_path):
            continue
        try:
            if "CJK" not in registered:
                pdfmetrics.registerFont(TTFont("CJK", reg_path))
                registered = pdfmetrics.getRegisteredFontNames()
        except Exception:
            continue  # 文件存在但 reportlab 无法加载(如 PostScript 轮廓的 ttc), 试下一个
        # 普通字体注册成功, 再尝试粗体
        for bold_path in _CJK_BOLD_FONTS:
            if not os.path.exists(bold_path):
                continue
            try:
                if "CJK-Bold" not in registered:
                    pdfmetrics.registerFont(TTFont("CJK-Bold", bold_path))
                return "CJK", "CJK-Bold"
            except Exception:
                continue
        return "CJK", "CJK"  # 无可用粗体时粗体退化为普通

    # 2. 回退 reportlab 内置中文 CID 字体(无需任何字体文件)
    try:
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        if "STSong-Light" not in registered:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light", "STSong-Light"
    except Exception:
        return "Helvetica", "Helvetica-Bold"  # 最终兜底


def render_pdf(model: DocumentModel, out_path: str):
    """将文档模型渲染为 PDF (版式重建)"""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    )
    from reportlab.lib.utils import simpleSplit

    font_name, font_bold = _register_cjk_fonts()

    styles = {
        "body": ParagraphStyle(
            "body", fontName=font_name, fontSize=11, leading=18,
            alignment=TA_JUSTIFY, wordWrap="CJK"),
        "heading1": ParagraphStyle(
            "h1", fontName=font_bold, fontSize=20, leading=26,
            spaceBefore=14, spaceAfter=10, textColor=colors.black),
        "heading2": ParagraphStyle(
            "h2", fontName=font_bold, fontSize=16, leading=22,
            spaceBefore=10, spaceAfter=8),
        "heading3": ParagraphStyle(
            "h3", fontName=font_bold, fontSize=13, leading=18,
            spaceBefore=8, spaceAfter=6),
        "list": ParagraphStyle(
            "list", fontName=font_name, fontSize=11, leading=18,
            leftIndent=18, bulletIndent=6, wordWrap="CJK"),
        "table": ParagraphStyle(
            "table", fontName=font_name, fontSize=10, leading=14, wordWrap="CJK"),
    }

    doc = SimpleDocTemplate(
        out_path, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=model.meta.get("title", "English Translator"),
    )
    story = []

    for b in model.blocks:
        text = b.text or ""
        if b.kind == KIND_EMPTY or not text.strip():
            story.append(Spacer(1, 6))
            continue
        if b.kind == KIND_HEADING:
            level = b.meta.get("level", 1)
            st = styles.get(f"heading{min(level, 3)}", styles["heading1"])
            story.append(Paragraph(_escape(text), st))
        elif b.kind == KIND_LIST_ITEM:
            story.append(Paragraph(_escape(text), styles["list"],
                                   bulletText="•"))
        else:
            align = b.style.get("align", 4)
            st = styles["body"]
            if align == 1:
                st = ParagraphStyle("c", parent=st, alignment=TA_CENTER)
            elif align == 2:
                st = ParagraphStyle("r", parent=st, alignment=TA_RIGHT)
            if b.style.get("bold"):
                st = ParagraphStyle("b", parent=st, fontName=font_bold)
            story.append(Paragraph(_escape(text), st))

    # 表格
    for table_model in model.tables:
        if not table_model:
            continue
        data = []
        for row in table_model:
            data.append([
                Paragraph(_escape(" ".join(
                    cb.text for cb in cell
                    if cb.text and cb.text.strip())), styles["table"])
                for cell in row
            ])
        if data:
            t = Table(data, hAlign="LEFT")
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(t)
            story.append(Spacer(1, 8))

    doc.build(story)


def render_pdf_bilingual(src_model: DocumentModel, tgt_model: DocumentModel,
                         out_path: str):
    """
    双语对照 PDF: 每个内容块先译文(正常)后原文(灰色小字), 表格译文在上原文在下。
    用于翻译后核对; 不改变 render_pdf 的单译文导出行为。
    """
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    )

    font_name, font_bold = _register_cjk_fonts()

    styles = {
        "body": ParagraphStyle(
            "body", fontName=font_name, fontSize=11, leading=18,
            alignment=TA_JUSTIFY, wordWrap="CJK"),
        "heading1": ParagraphStyle(
            "h1", fontName=font_bold, fontSize=20, leading=26,
            spaceBefore=14, spaceAfter=10, textColor=colors.black),
        "heading2": ParagraphStyle(
            "h2", fontName=font_bold, fontSize=16, leading=22,
            spaceBefore=10, spaceAfter=8),
        "heading3": ParagraphStyle(
            "h3", fontName=font_bold, fontSize=13, leading=18,
            spaceBefore=8, spaceAfter=6),
        "list": ParagraphStyle(
            "list", fontName=font_name, fontSize=11, leading=18,
            leftIndent=18, bulletIndent=6, wordWrap="CJK"),
        "table": ParagraphStyle(
            "table", fontName=font_name, fontSize=10, leading=14, wordWrap="CJK"),
        # 原文样式: 灰色小字, 与译文区分
        "src": ParagraphStyle(
            "src", fontName=font_name, fontSize=8.5, leading=13,
            textColor=colors.grey, wordWrap="CJK"),
        "src_table": ParagraphStyle(
            "src_table", fontName=font_name, fontSize=8, leading=12,
            textColor=colors.grey, wordWrap="CJK"),
    }

    doc = SimpleDocTemplate(
        out_path, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=tgt_model.meta.get("title", "Bilingual Translation"),
    )
    story = []

    def _render_block(b, is_src: bool):
        """渲染单个文本块; is_src=True 时用灰色小字样式"""
        text = b.text or ""
        if not text.strip():
            story.append(Spacer(1, 4))
            return
        if b.kind == KIND_HEADING:
            level = b.meta.get("level", 1)
            st = styles.get(f"heading{min(level, 3)}", styles["heading1"])
            if is_src:
                st = ParagraphStyle(
                    "src-h", parent=st, fontName=font_name, fontSize=9,
                    leading=14, textColor=colors.grey)
            story.append(Paragraph(_escape(text), st))
        elif b.kind == KIND_LIST_ITEM:
            st = styles["src"] if is_src else styles["list"]
            story.append(Paragraph(_escape(text), st, bulletText="•"))
        else:
            align = b.style.get("align", 4)
            st = styles["src"] if is_src else styles["body"]
            if align == 1:
                st = ParagraphStyle("c", parent=st, alignment=TA_CENTER)
            elif align == 2:
                st = ParagraphStyle("r", parent=st, alignment=TA_RIGHT)
            if b.style.get("bold") and not is_src:
                st = ParagraphStyle("b", parent=st, fontName=font_bold)
            story.append(Paragraph(_escape(text), st))

    def _render_table(tbl, is_src: bool):
        if not tbl:
            return
        data = []
        for row in tbl:
            data.append([
                Paragraph(_escape(" ".join(
                    cb.text for cb in cell
                    if cb.text and cb.text.strip())),
                    styles["src_table"] if is_src else styles["table"])
                for cell in row
            ])
        if data:
            t = Table(data, hAlign="LEFT")
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(t)

    src_blocks = list(src_model.blocks) if src_model else []

    # 内容块: 译文 + 原文(上下对照)
    for i, b in enumerate(tgt_model.blocks):
        if b.kind == KIND_EMPTY or not (b.text or "").strip():
            story.append(Spacer(1, 6))
            continue
        _render_block(b, is_src=False)          # 译文
        s_b = src_blocks[i] if i < len(src_blocks) else None
        if s_b and (s_b.text or "").strip():
            _render_block(s_b, is_src=True)     # 原文(灰色小字)
        story.append(Spacer(1, 8))              # 对照块之间留间隔

    # 表格: 译文表格 + 原文表格
    for ti, tbl in enumerate(tgt_model.tables):
        if not tbl:
            continue
        _render_table(tbl, is_src=False)
        src_tbl = (src_model.tables[ti]
                   if src_model and ti < len(src_model.tables) else None)
        if src_tbl:
            _render_table(src_tbl, is_src=True)
        story.append(Spacer(1, 8))

    doc.build(story)


def _escape(text: str) -> str:
    """转义 XML 特殊字符(Paragraph 使用 XML 标记)"""
    return (text.replace("&", "&amp;").replace("<", "&lt;")
                .replace(">", "&gt;"))


# ---------------------------------------------------------------
# 统一入口
# ---------------------------------------------------------------
def export_document(model: DocumentModel, fmt: str, out_path: str):
    """按格式导出文档; fmt: docx | pdf"""
    if fmt == "docx":
        render_docx(model, out_path)
    elif fmt == "pdf":
        render_pdf(model, out_path)
    else:
        raise ValueError(f"不支持的导出格式: {fmt}")
