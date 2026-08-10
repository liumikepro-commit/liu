# -*- coding: utf-8 -*-
"""
parser.py — 文档解析模块
支持: PDF (.pdf), Word (.docx)
- PDF : pypdf 提取文本层, 按空行分块
- DOCX: python-docx 完整解析段落/标题/列表/表格, 保留格式元数据
- .doc: 老格式二进制, 无法可靠解析, 给出清晰提示(建议另存为 .docx)

统一输出: DocumentModel (见 block.py)
"""
import os
import re
from typing import Union

from .block import DocumentModel, TextBlock

# 支持格式白名单(小写扩展名 -> 描述)
SUPPORTED_FORMATS = {
    "pdf": "PDF 文档",
    "docx": "Word 文档 (.docx)",
}
# 明确识别但无法解析的格式
UNSUPPORTED_FORMATS = {
    "doc": "旧版 Word (.doc) 为二进制格式，请先用 Word 另存为 .docx 后再上传。",
    "txt": "纯文本文件请使用上方的「文本翻译」功能。",
    "xls": "Excel 文件 (.xls/.xlsx) 不在支持范围内，当前仅支持 PDF 与 Word。",
    "xlsx": "Excel 文件 (.xls/.xlsx) 不在支持范围内，当前仅支持 PDF 与 Word。",
    "ppt": "PowerPoint 文件 (.ppt/.pptx) 不在支持范围内，当前仅支持 PDF 与 Word。",
    "pptx": "PowerPoint 文件 (.ppt/.pptx) 不在支持范围内，当前仅支持 PDF 与 Word。",
}


def detect_format(filename: str) -> str:
    """根据文件名返回小写扩展名"""
    ext = os.path.splitext(filename)[1].lower().lstrip(".")
    return ext


def validate_format(filename: str) -> str:
    """
    校验文件格式, 返回扩展名; 不支持时抛出 ValueError(带清晰中文提示)。
    """
    ext = detect_format(filename)
    if not ext:
        raise ValueError("无法识别文件格式，请上传 .pdf 或 .docx 文件。")
    if ext in SUPPORTED_FORMATS:
        return ext
    if ext in UNSUPPORTED_FORMATS:
        raise ValueError(UNSUPPORTED_FORMATS[ext])
    raise ValueError(
        f"不支持的文件格式 .{ext}，当前支持 PDF (.pdf) 与 Word (.docx)。"
    )


# ---------------------------------------------------------------
# PDF 解析
# ---------------------------------------------------------------
def parse_pdf(path: str) -> DocumentModel:
    """
    解析 PDF: 提取文本层, 按空行分段为 TextBlock。
    说明: PDF 无固定版式结构, 采用"文本流 + 空行分段"的提取策略;
    标题/字号等样式信息无法从文本层可靠获取, 重建时按普通段落处理。
    """
    from pypdf import PdfReader

    try:
        reader = PdfReader(path)
    except Exception as e:
        raise ValueError(f"PDF 解析失败（文件可能已损坏或加密）：{e}")

    if reader.is_encrypted:
        raise ValueError("PDF 已加密，请先解除密码保护后再上传。")

    model = DocumentModel()
    model.meta["source_format"] = "pdf"
    model.meta["page_count"] = len(reader.pages)

    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if not text.strip():
            continue
        # 按空行分段, 每段为一个块
        for para in text.split("\n\n"):
            para = para.strip()
            if not para:
                continue
            # 块内按换行合并为连续文本(保留语义), 还原 PDF 行断
            merged = re.sub(r"\s*\n\s*", " ", para)
            model.blocks.append(TextBlock(text=merged, kind="paragraph"))

    if not model.blocks:
        raise ValueError(
            "未能从 PDF 中提取到文本。该 PDF 可能是扫描件（图片型），"
            "请先使用 OCR 工具识别为文本型 PDF 后再上传。"
        )
    return model


# ---------------------------------------------------------------
# DOCX 解析
# ---------------------------------------------------------------
def _extract_para_style(para) -> dict:
    """提取段落级格式元数据(对齐/粗体/斜体/字号/颜色)"""
    style = {}
    try:
        if para.alignment is not None:
            style["align"] = int(para.alignment)
    except Exception:
        pass
    # 汇总段内 run 格式: 取多数 run 的格式
    runs = [r for r in para.runs if r.text.strip()]
    if runs:
        bold_flags = [bool(r.bold) for r in runs]
        italic_flags = [bool(r.italic) for r in runs]
        style["bold"] = sum(bold_flags) > len(bold_flags) / 2
        style["italic"] = sum(italic_flags) > len(italic_flags) / 2
        sizes = [r.font.size for r in runs if r.font.size]
        if sizes:
            style["font_size"] = max(int(s.pt) for s in sizes)
    return style


def _heading_level(style_name: str) -> int:
    """从 Word 样式名提取标题级别, 返回 0 表示非标题"""
    m = re.search(r"Heading\s*(\d)|标题\s*(\d)", style_name or "", re.I)
    if m:
        return int(m.group(1) or m.group(2))
    return 0


def parse_docx(path: str) -> DocumentModel:
    """解析 .docx: 段落/标题/列表/表格, 保留格式元数据"""
    from docx import Document

    try:
        doc = Document(path)
    except Exception as e:
        raise ValueError(f"Word 文档解析失败（文件可能已损坏）：{e}")

    model = DocumentModel()
    model.meta["source_format"] = "docx"

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        level = _heading_level(style_name)

        if not text:
            model.blocks.append(TextBlock(text="", kind="empty"))
            continue

        style = _extract_para_style(para)
        if level > 0:
            block = TextBlock(text=text, kind="heading",
                              style=style, meta={"level": min(level, 6)})
        elif style_name and ("List" in style_name or "列表" in style_name):
            block = TextBlock(text=text, kind="list_item",
                              style=style, meta={"level": level or 1})
        else:
            block = TextBlock(text=text, kind="paragraph", style=style)
        model.blocks.append(block)

    # 表格
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_blocks = [
                    TextBlock(text=p.text.strip(), kind="paragraph",
                              style=_extract_para_style(p))
                    for p in cell.paragraphs if p.text.strip()
                ]
                if not cell_blocks:
                    cell_blocks = [TextBlock(text="", kind="empty")]
                cells.append(cell_blocks)
            rows.append(cells)
        model.tables.append(rows)

    return model


# ---------------------------------------------------------------
# 统一入口
# ---------------------------------------------------------------
def parse_document(filename: str, file_path: str) -> DocumentModel:
    """
    解析入口: 根据扩展名分派到对应解析器。
    返回 DocumentModel; 解析失败抛 ValueError(中文提示)。
    """
    fmt = validate_format(filename)
    if fmt == "pdf":
        return parse_pdf(file_path)
    if fmt == "docx":
        return parse_docx(file_path)
    raise ValueError(f"不支持的格式: {fmt}")  # 理论上不会到达
