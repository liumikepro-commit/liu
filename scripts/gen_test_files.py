# -*- coding: utf-8 -*-
"""生成英文测试文件: test_english.docx + test_english.txt"""
from docx import Document

doc = Document()
doc.add_heading('Annual Report of TechNova Corporation', level=0)
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'TechNova Corporation is a global leader in cloud computing and artificial '
    'intelligence solutions. In fiscal year 2025, the company achieved record '
    'revenue of 12.8 billion US dollars, representing a year-over-year growth '
    'of 18 percent. Our mission is to empower every organization on the planet '
    'to achieve more through intelligent technology.'
)
doc.add_paragraph(
    'This report provides an overview of our business performance, strategic '
    'initiatives, and financial highlights for the past twelve months.'
)

doc.add_heading('Business Performance', level=1)
doc.add_paragraph(
    'Our cloud services division continued to be the primary growth engine, '
    'contributing 62 percent of total revenue. The number of enterprise '
    'customers increased by 25 percent, and customer retention rate reached '
    'a record high of 94 percent.'
)

doc.add_heading('Key Financial Highlights', level=1)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
data = [
    ('Metric', 'FY 2024', 'FY 2025'),
    ('Revenue (USD Billion)', '10.8', '12.8'),
    ('Net Income (USD Billion)', '1.9', '2.4'),
    ('R&D Investment (USD Billion)', '1.2', '1.5'),
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = table.cell(i, j)
        cell.text = val
        if i == 0:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_heading('Strategic Initiatives', level=1)
doc.add_paragraph('We invested heavily in three strategic areas this year:', style='List Bullet')
doc.add_paragraph('Expansion of our global data center network to support AI workloads', style='List Bullet')
doc.add_paragraph('Development of next-generation large language models and translation services', style='List Bullet')
doc.add_paragraph('Partnership with major universities to advance fundamental AI research', style='List Bullet')

doc.add_heading('Outlook', level=1)
doc.add_paragraph(
    'Looking ahead to fiscal year 2026, we expect continued strong growth '
    'driven by demand for artificial intelligence infrastructure and '
    'enterprise digital transformation. We are confident in our ability to '
    'deliver sustainable value to our shareholders, customers, and employees.'
)
doc.add_paragraph(
    'For more information about TechNova Corporation and its products, '
    'please visit our website at www.technova.example.com or contact our '
    'investor relations team.'
)

doc.save('uploads/test_english.docx')
print('docx 生成成功')

text = (
    "Annual Report of TechNova Corporation\n\n"
    "Executive Summary\n"
    "TechNova Corporation is a global leader in cloud computing and artificial intelligence solutions. "
    "In fiscal year 2025, the company achieved record revenue of 12.8 billion US dollars, "
    "representing a year-over-year growth of 18 percent.\n\n"
    "Business Performance\n"
    "Our cloud services division continued to be the primary growth engine, contributing 62 percent of total revenue. "
    "The number of enterprise customers increased by 25 percent, and customer retention rate reached a record high of 94 percent.\n\n"
    "Strategic Initiatives\n"
    "We invested heavily in three strategic areas this year. First, we expanded our global data center network to support AI workloads. "
    "Second, we developed next-generation large language models and translation services. "
    "Third, we formed partnerships with major universities to advance fundamental AI research.\n\n"
    "Outlook\n"
    "Looking ahead to fiscal year 2026, we expect continued strong growth driven by demand for artificial intelligence infrastructure "
    "and enterprise digital transformation."
)
with open('uploads/test_english.txt', 'w', encoding='utf-8') as f:
    f.write(text)
print('txt 生成成功,', len(text), '字符')
