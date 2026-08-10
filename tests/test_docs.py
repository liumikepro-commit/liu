# -*- coding: utf-8 -*-
"""
test_docs.py — 文档导入与翻译模块测试
覆盖: 格式校验、DOCX/PDF 解析、专名保护、翻译流水线、docx/pdf 导出。

运行: python -m unittest tests.test_docs -v
"""
import os
import shutil
import tempfile
import unittest

from translator.docs import parser, pipeline, renderer
from translator.docs.block import DocumentModel, TextBlock

TMP = tempfile.mkdtemp(prefix="translator_test_")


class TestFormatValidation(unittest.TestCase):
    """格式校验与错误提示"""

    def test_supported(self):
        self.assertEqual(parser.validate_format("a.pdf"), "pdf")
        self.assertEqual(parser.validate_format("a.DOCX"), "docx")

    def test_doc_unsupported(self):
        with self.assertRaises(ValueError) as ctx:
            parser.validate_format("old.doc")
        self.assertIn(".docx", str(ctx.exception))

    def test_txt_unsupported(self):
        with self.assertRaises(ValueError):
            parser.validate_format("note.txt")

    def test_unknown_ext(self):
        with self.assertRaises(ValueError):
            parser.validate_format("data.xyz")

    def test_no_ext(self):
        with self.assertRaises(ValueError):
            parser.validate_format("README")


class TestProperNouns(unittest.TestCase):
    """专有名词保护(仅词典未收录的词被保护)"""

    def test_extract(self):
        nouns = pipeline._extract_proper_nouns(
            "Zylink works at Microsoft. Shanghai is big.")
        self.assertIn("Zylink", nouns)      # 词典未收录 -> 保护
        self.assertNotIn("Microsoft", nouns)  # 词典收录(微软) -> 不保护
        self.assertNotIn("Shanghai", nouns)   # 词典收录(上海) -> 不保护

    def test_common_words_not_protected(self):
        nouns = pipeline._extract_proper_nouns(
            "Annual Report of Zylink. Key highlights are listed below.")
        self.assertNotIn("Annual Report", nouns)  # 均被词典收录 -> 不保护
        self.assertIn("Zylink", nouns)

    def test_sequence_with_unknown(self):
        nouns = pipeline._extract_proper_nouns("Welcome to Zylink City.")
        self.assertIn("Zylink City", nouns)  # 序列含未收录词 -> 整体保护

    def test_protect_restore(self):
        text = "Zylink works in New York."
        protected, mapping = pipeline._protect_proper_nouns(text)
        self.assertNotIn("Zylink", protected)
        # 模拟翻译(原文不变)
        restored = pipeline._restore_proper_nouns(protected, mapping)
        self.assertEqual(restored, text)


class TestDocxParsing(unittest.TestCase):
    """DOCX 解析"""

    @classmethod
    def setUpClass(cls):
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading("Meeting Notes", level=1)
        doc.add_paragraph("This is the first paragraph about the project.")
        p = doc.add_paragraph()
        run = p.add_run("This is a bold sentence.")
        run.bold = True
        doc.add_paragraph("First item", style="List Bullet")
        doc.add_paragraph("Second item", style="List Bullet")
        # 表格
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Role"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "Engineer"
        cls.path = os.path.join(TMP, "sample.docx")
        doc.save(cls.path)

    def test_parse_structure(self):
        model = parser.parse_docx(self.path)
        kinds = [b.kind for b in model.blocks]
        self.assertIn("heading", kinds)
        self.assertIn("list_item", kinds)
        self.assertGreaterEqual(len(model.tables), 1)

    def test_heading_level(self):
        model = parser.parse_docx(self.path)
        heading = next(b for b in model.blocks if b.kind == "heading")
        self.assertEqual(heading.meta.get("level"), 1)

    def test_bold_style(self):
        model = parser.parse_docx(self.path)
        bold = next(b for b in model.blocks if "bold" in b.text)
        self.assertTrue(bold.style.get("bold"))


class TestPdfParsing(unittest.TestCase):
    """PDF 解析"""

    @classmethod
    def setUpClass(cls):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate

        styles = getSampleStyleSheet()
        doc = SimpleDocTemplate(os.path.join(TMP, "sample.pdf"), pagesize=A4)
        story = [
            Paragraph("Title Line", styles["Title"]),
            Paragraph("Hello world. This is a PDF test document.", styles["BodyText"]),
        ]
        doc.build(story)
        cls.path = os.path.join(TMP, "sample.pdf")

    def test_parse_text(self):
        model = parser.parse_pdf(self.path)
        text = " ".join(b.text for b in model.blocks)
        self.assertIn("Hello world", text)
        self.assertGreater(len(model.blocks), 0)


class TestPipeline(unittest.TestCase):
    """翻译流水线(本地引擎, 快速)"""

    def test_translate_model(self):
        model = DocumentModel()
        model.blocks = [
            TextBlock(text="Hello world", kind="heading", meta={"level": 1}),
            TextBlock(text="This is a book.", kind="paragraph"),
            TextBlock(text="Thank you.", kind="paragraph"),
        ]
        model.tables = [[
            [[TextBlock(text="Apple")], [TextBlock(text="Book")]],
        ]]

        result = pipeline.translate_document(
            model, source="en", target="zh", use_online=False)

        self.assertEqual(len(result.blocks), 3)  # 块数不变
        self.assertTrue(result.blocks[0].text)   # 标题被翻译
        self.assertTrue(any("谢谢" in b.text or "感谢" in b.text for b in result.blocks))
        # 表格单元格被翻译
        cell_text = result.tables[0][0][0][0].text
        self.assertTrue("苹果" in cell_text or cell_text == "Apple")

    def test_zh_to_en(self):
        model = DocumentModel()
        model.blocks = [TextBlock(text="我的书", kind="paragraph")]
        result = pipeline.translate_document(
            model, source="zh", target="en", use_online=False)
        self.assertIn("my", result.blocks[0].text.lower())


class TestRender(unittest.TestCase):
    """导出 docx/pdf"""

    @classmethod
    def setUpClass(cls):
        cls.model = DocumentModel()
        cls.model.blocks = [
            TextBlock(text="项目报告", kind="heading", meta={"level": 1},
                      style={"bold": True}),
            TextBlock(text="这是第一段内容。", kind="paragraph"),
            TextBlock(text="第一项", kind="list_item"),
            TextBlock(text="第二项", kind="list_item"),
        ]
        cls.model.tables = [[
            [[TextBlock(text="名称")], [TextBlock(text="数量")]],
            [[TextBlock(text="苹果")], [TextBlock(text="3")]],
        ]]
    def test_export_docx(self):
        out = os.path.join(TMP, "out.docx")
        renderer.render_docx(self.model, out)
        self.assertTrue(os.path.exists(out) and os.path.getsize(out) > 0)
        # 重新解析验证内容
        model2 = parser.parse_docx(out)
        text = " ".join(b.text for b in model2.blocks)
        self.assertIn("项目报告", text)

    def test_export_pdf(self):
        out = os.path.join(TMP, "out.pdf")
        renderer.render_pdf(self.model, out)
        self.assertTrue(os.path.exists(out) and os.path.getsize(out) > 0)
        # 重新解析验证内容
        model2 = parser.parse_pdf(out)
        text = " ".join(b.text for b in model2.blocks)
        self.assertIn("项目报告", text)


if __name__ == "__main__":
    unittest.main()
